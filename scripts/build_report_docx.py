#!/usr/bin/env python3
"""Build the Archivist evaluation report as a polished Word document."""

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "debugging-evaluation-report.md"
OUTPUT = ROOT / "docs" / "archivist-model-testing-debugging-report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E0E8"


def set_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    set_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, label, target):
    relationship = paragraph.part.relate_to(
        target,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([props, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|_[^_]+_|\[[^]]+\]\([^)]+\))")


def add_inline(paragraph, text):
    cursor = 0
    for match in INLINE.finditer(text):
        if match.start() > cursor:
            set_font(paragraph.add_run(text[cursor:match.start()]))
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=9.5, name="IBM Plex Mono")
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, bold=True)
        elif token.startswith("_"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, italic=True)
        else:
            label, target = re.match(r"\[([^]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, target)
        cursor = match.end()
    if cursor < len(text):
        set_font(paragraph.add_run(text[cursor:]))


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_header_row(row):
    row_props = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    row_props.append(header)


def shade_cell(cell, fill):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_results_table(doc, lines):
    rows = [[part.strip() for part in line.strip().strip("|").split("|")] for line in lines]
    rows = [rows[0]] + rows[2:]
    table = doc.add_table(rows=len(rows), cols=4)
    table.style = "Table Grid"
    # Named override: seven-inch table for the three-page academic submission.
    widths = [3200, 1900, 1900, 3080]
    for r_index, row in enumerate(rows):
        for c_index, value in enumerate(row):
            paragraph = table.cell(r_index, c_index).paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            set_font(run, size=8.5, bold=(r_index == 0))
            if r_index == 0:
                shade_cell(table.cell(r_index, c_index), LIGHT_GRAY)
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    # Named override: compact academic submission, targeting the assigned
    # three-page limit while preserving the full analysis.
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 14, BLUE, 10, 4),
        ("Heading 2", 12, BLUE, 8, 4),
        ("Heading 3", 10.5, DARK_BLUE, 6, 3),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(9.5)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(2)
    bullet.paragraph_format.line_spacing = 1.0


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "ARCHIVIST  |  MODEL EVALUATION"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_font(header.runs[0], size=8.5, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Archivist model testing and debugging report")
    set_font(run, size=20, bold=True, color=RGBColor(11, 37, 69))
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("A/B evaluation of direct generation and retrieval-augmented generation")
    set_font(run, size=10.5, italic=True, color=MUTED)

    metadata = doc.add_paragraph()
    metadata.paragraph_format.space_after = Pt(12)
    add_inline(metadata, "Dataset: `openstax-1.3-v1`  |  Model: `gemma3:1b`  |  Questions: 15  |  Conditions: 2")

    lines = SOURCE.read_text().splitlines()[1:]
    index = 0
    in_code = False
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            in_code = not in_code
            index += 1
            continue
        if in_code:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(line)
            set_font(run, size=9, name="IBM Plex Mono", color=RGBColor(31, 41, 55))
            index += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("| "):
            table_lines = []
            while index < len(lines) and lines[index].startswith("| "):
                table_lines.append(lines[index])
                index += 1
            add_results_table(doc, table_lines)
            continue
        elif line.startswith("- "):
            bullet_text = line[2:].strip()
            while index + 1 < len(lines):
                following = lines[index + 1]
                if not following.strip() or following.startswith(("- ", "#", "|", "```")):
                    break
                bullet_text += " " + following.strip()
                index += 1
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline(paragraph, bullet_text)
        elif line.strip():
            paragraph_text = line.strip()
            while index + 1 < len(lines):
                following = lines[index + 1]
                if not following.strip() or following.startswith(("- ", "#", "|", "```")):
                    break
                paragraph_text += " " + following.strip()
                index += 1
            paragraph = doc.add_paragraph()
            add_inline(paragraph, paragraph_text)
        index += 1

    properties = doc.core_properties
    properties.title = "Archivist model testing and debugging report"
    properties.subject = "A/B evaluation of direct and retrieval-augmented generation"
    properties.author = "Nick Moore"
    properties.keywords = "Archivist, RAG, evaluation, debugging, Ollama"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
