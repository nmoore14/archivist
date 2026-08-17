from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path('/Users/nickmoore/Desktop/school/summer_2026/msai_699/archivist')
OUT = ROOT / 'Archivist_Final_Technical_Report.docx'
BLUE = '1F4D78'; ACCENT = '2F80ED'; NAVY = '173B57'; MUTED = '5B6570'; LIGHT = 'E8EEF5'; PALE = 'F4F6F9'; GOLD = 'D97706'; WHITE='FFFFFF'; BLACK='17191C'

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
# Named override: academic_page_budget. Keeps the assigned report within seven
# pages while preserving readable type and the narrative-proposal hierarchy.
sec.top_margin = sec.bottom_margin = Inches(0.72)
sec.left_margin = sec.right_margin = Inches(0.8)
sec.header_distance = sec.footer_distance = Inches(0.492)
sec.different_first_page_header_footer = True

def set_font(run, size=11, bold=False, italic=False, color=BLACK, name='Calibri'):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None: node = OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'),'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); tblHeader = OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)

def set_table_geometry(table, widths_dxa):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW')); tblW.set(qn('w:w'), str(sum(widths_dxa))); tblW.set(qn('w:type'),'dxa')
    tblInd = OxmlElement('w:tblInd'); tblInd.set(qn('w:w'),'120'); tblInd.set(qn('w:type'),'dxa'); tblPr.append(tblInd)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths_dxa:
        col=OxmlElement('w:gridCol'); col.set(qn('w:w'),str(w)); grid.append(col)
    for row in table.rows:
        for i,cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[i]/1440)
            tcW=cell._tc.get_or_add_tcPr().find(qn('w:tcW')); tcW.set(qn('w:w'),str(widths_dxa[i])); tcW.set(qn('w:type'),'dxa')
            set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r=paragraph.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); r._r.addnext(fld)

# Explicit narrative_proposal tokens
normal = doc.styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(BLACK)
normal.paragraph_format.space_before=Pt(0); normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.15
for style_name,size,color,before,after in [('Heading 1',16,BLUE,12,7),('Heading 2',13,BLUE,8,4),('Heading 3',12,NAVY,6,3)]:
    st=doc.styles[style_name]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
caption = doc.styles['Caption']; caption.font.name='Calibri'; caption.font.size=Pt(9); caption.font.italic=True; caption.font.color.rgb=RGBColor.from_string(MUTED)
caption.paragraph_format.space_before=Pt(4); caption.paragraph_format.space_after=Pt(8); caption.paragraph_format.keep_with_next=False
if 'Report Lead' not in doc.styles:
    lead=doc.styles.add_style('Report Lead',WD_STYLE_TYPE.PARAGRAPH); lead.base_style=normal; lead.font.name='Calibri'; lead.font.size=Pt(13); lead.font.color.rgb=RGBColor.from_string(NAVY)
    lead.paragraph_format.space_after=Pt(9); lead.paragraph_format.line_spacing=1.15

# running furniture
hp=sec.header.paragraphs[0]; hp.text='ARCHIVIST  |  FINAL TECHNICAL REPORT'; hp.alignment=WD_ALIGN_PARAGRAPH.LEFT
for r in hp.runs: set_font(r,9,True,False,MUTED)
fp=sec.footer.paragraphs[0]; add_page_number(fp)
for r in fp.runs: set_font(r,9,False,False,MUTED)

def add_p(text='', style=None, bold_start=None, align=None, before=0, after=None, keep=False):
    p=doc.add_paragraph(style=style)
    if bold_start and text.startswith(bold_start):
        r=p.add_run(bold_start); set_font(r,bold=True)
        r=p.add_run(text[len(bold_start):]); set_font(r)
    else:
        r=p.add_run(text); set_font(r, size=13 if style=='Report Lead' else 11, color=NAVY if style=='Report Lead' else BLACK)
    if align is not None: p.alignment=align
    p.paragraph_format.space_before=Pt(before)
    if after is not None: p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.keep_with_next=keep
    return p

def heading(text, level=1): return doc.add_heading(text, level=level)
def page_break(): doc.add_page_break()

def callout(label, value, detail):
    t=doc.add_table(rows=1,cols=2); set_table_geometry(t,[2160,7200])
    set_repeat_table_header(t.rows[0])
    set_cell_shading(t.cell(0,0),NAVY); set_cell_shading(t.cell(0,1),PALE)
    p=t.cell(0,0).paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(label); set_font(r,9,True,False,WHITE)
    p=t.cell(0,1).paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(value+'  '); set_font(r,14,True,False,ACCENT)
    r=p.add_run(detail); set_font(r,10,False,False,MUTED)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def data_table(headers, rows, widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; set_table_geometry(t,widths); set_repeat_table_header(t.rows[0])
    for i,h in enumerate(headers):
        set_cell_shading(t.cell(0,i),LIGHT); p=t.cell(0,i).paragraphs[0]; p.paragraph_format.space_after=Pt(0)
        r=p.add_run(h); set_font(r,9,True,False,NAVY)
        if i: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0)
            r=p.add_run(str(val)); set_font(r,9,False,False,BLACK)
            if i: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(t,widths)
    return t

# PAGE 1 — editorial cover + abstract
for _ in range(3): add_p('',after=12)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('MSAI 699 CAPSTONE PROJECT'); set_font(r,11,True,False,ACCENT)
p.paragraph_format.space_after=Pt(20)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Archivist'); set_font(r,30,True,False,NAVY)
p.paragraph_format.space_after=Pt(6)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Investigating Guided Workflows for Reducing Technical Barriers to Local Artificial Intelligence Adoption'); set_font(r,16,False,False,BLUE)
p.paragraph_format.space_after=Pt(18)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Nick Moore  |  Summer 2026'); set_font(r,11,False,False,MUTED)
p.paragraph_format.space_after=Pt(34)
heading('Abstract',1)
add_p('Archivist is a local-first artificial intelligence workspace for schools. Administrators create course workspaces and upload trusted materials; assigned students ask questions and receive locally generated answers with source references. The project investigates whether guided workflows can reduce the technical barriers associated with configuring and deploying local AI assistants for educational use. A reproducible evaluation compared direct generation with retrieval-augmented generation (RAG) across 15 versioned questions, followed by a six-configuration RAG optimization study. RAG increased manually scored correctness from 63.3% to 93.3%. The selected top-k-one configuration achieved 96.7% correctness, 0.97 average groundedness, a 100% retrieval hit rate, and a 1.07-second mean response time. These results validate the technical value of grounding and tuning on the project test set, while the central usability claim still requires stakeholder testing. Future work should evaluate administrator task completion and confidence, then harden OCR, background indexing, security, and storage for broader deployment.')

# PAGE 2 — introduction
page_break(); heading('1. Introduction',1)
add_p('Educational institutions want the accessibility of conversational AI without surrendering control over course content, student interactions, and institutional data. Hosted assistants lower the barrier to experimentation, but they can create concerns involving privacy, source authority, recurring cost, network dependence, and unsupported answers. Local models address part of that problem by keeping inference and storage under institutional control; however, they introduce a different barrier: setup and maintenance commonly require familiarity with containers, model servers, embeddings, retrieval configuration, and networking.', style='Report Lead')
add_p('The central research question is whether guided workflows can significantly reduce the technical barriers associated with configuring and deploying local AI assistants for educational use. Archivist makes this question concrete through an end-to-end minimum viable product (MVP). An administrator can create a workspace, upload approved course sources, add and assign students, and expose a grounded question-answering experience. The system is designed for a school or small organization that values privacy and control but may not have a dedicated AI infrastructure team.')
add_p('The work pursued three connected objectives. First, it built a usable local application rather than an isolated notebook demonstration. Second, it tested whether retrieval improves source-specific answers over direct local generation. Third, it optimized retrieval settings to identify a practical quality–latency balance. The broader research objective—whether guided deployment reduces human effort and errors—remains the next empirical phase.')
add_p('Technically, Archivist uses a Go HTTP server, server-rendered HTML with partial updates, SQLite storage, and Ollama for local chat and embedding inference. The application separates interface handlers from shared services. Uploaded files are extracted, divided into overlapping chunks, embedded, and stored with pipeline-version metadata. For each student question, the retrieval service embeds the query, ranks stored chunks using cosine similarity, supplies selected evidence to a context-constrained prompt, and returns an answer with source references. Docker Compose packages the application and model service into a repeatable local deployment.')
callout('DESIGN PRINCIPLE','Local by default','Course documents, embeddings, model calls, accounts, and chat history remain under the operator’s control.')

# PAGE 3 — methodology
page_break(); heading('2. Methodology',1)
add_p('Development followed an iterative workflow: define the stakeholder path, implement the smallest complete vertical slice, test it locally, document failure modes, and convert those failures into regression checks. The resulting path begins with first-run administration and continues through workspace creation, source ingestion, assignment, retrieval, and cited response generation. The web interface and terminal interface share application services so business rules are not duplicated in templates or handlers.')
add_p('For the baseline experiment, the evaluation used 15 versioned questions about a bundled, instructor-prepared paraphrase of an OpenStax data-science section. The questions covered definitions, comparisons, recall, reasoning, synthesis, and two deliberately unanswerable items. Both conditions used the same local chat model and deterministic generation settings. The no-RAG condition received only the question. The RAG condition additionally received the three most similar source chunks, using 1,000-character chunks, 200-character overlap, nomic-embed-text embeddings, and cosine similarity.')
add_p('Each run captured responses, retrieved chunk identifiers, similarity scores, retrieval time, generation time, and total response time. Manual review scored correctness on a 0–2 rubric: 2 for correct and complete, 1 for partially correct or incomplete, and 0 for incorrect or missing. RAG answers were also scored for groundedness against the retrieved text. Retrieval hit rate measured whether retrieved context contained the evidence required by the reference answer, and hallucination handling measured whether unanswerable questions avoided unsupported content.')
add_p('The optimization experiment evaluated six RAG configurations using the same 15 questions. The study varied chunk size and overlap—600/100, 800/150, 1,000/200, and 1,400/250 characters—and varied top-k at 1, 3, and 5 for the 1,000/200 setting. A project-specific selection score combined correctness, groundedness, retrieval success, and response time. The score was used for engineering selection rather than as a general benchmark.')
add_p('To support reproducibility, the dataset, source text, experiment settings, run metadata, raw answers, manual-review files, and summaries were retained in the repository. Source and configuration hashes support comparison across runs. The design intentionally favors auditability, but it does not provide statistical significance: the question set is small, the source is narrow, and the manual rubric is reviewer-dependent.')

# PAGE 4 — baseline results
page_break(); heading('3. Results',1)
add_p('On the final baseline evaluation, direct generation achieved 63.3% of the maximum correctness score. Adding retrieved course context increased correctness to 93.3%, a gain of 30.0 percentage points. Retrieval Hit Rate@3 was 100%, average groundedness was 0.93 on a 0–1 scale, and hallucination handling on the questions marked unanswerable from the section was 100%. Mean response time increased modestly from 1.11 seconds without RAG to 1.25 seconds with RAG.')
data_table(['Metric','No RAG','RAG','Difference'],[
    ('Correctness','63.3%','93.3%','+30.0 pts'),
    ('Mean response time','1.11 s','1.25 s','+0.14 s'),
    ('Retrieval Hit Rate@3','—','100%','—'),
    ('Average groundedness','—','0.93','—'),
    ('Hallucination handling','—','100%','—'),
],[3600,1920,1920,1920])
p=doc.add_paragraph('Table 1. Final baseline comparison across 15 manually scored questions.',style='Caption')
add_p('The largest difference appeared when questions required source-specific wording or boundaries. Direct generation could produce plausible general knowledge while missing the assigned material’s intended explanation. On an out-of-scope question about transformer embeddings, the no-RAG answer invented a technically plausible explanation even though the section did not contain one; RAG correctly declined to answer from the available material. This behavior matters in education because a fluent answer can still conflict with the instructor-approved source.')
add_p('Retrieval success did not guarantee complete generation. Some answers used the correct evidence but omitted a required qualification or overstated the source. This distinction justified separate retrieval-hit, groundedness, and correctness measures rather than collapsing all behavior into one pass rate.')
callout('KEY RESULT','+30.0 points','Grounding improved correctness far more than the observed 0.14-second mean latency cost on this test set.')

# PAGE 5 — optimization and challenges
page_break()
add_p('All six configurations achieved a 100% retrieval hit rate and 100% hallucination handling on the small test set. Correctness ranged from 93.3% to 100%, while mean response time ranged from 1.07 to 1.88 seconds. The top-k-one configuration used 1,000-character chunks with 200-character overlap and retrieved one chunk. It achieved 96.7% correctness, 0.97 groundedness, and a 1.07-second mean response time. Relative to the top-k-three baseline, correctness increased 3.3 points and mean latency fell 0.81 seconds, or 42.9%.')
pic_run = doc.add_paragraph().add_run()
pic_run.add_picture(str(ROOT/'results/charts/correctness_latency_tradeoff.png'),width=Inches(5.25))
p=doc.paragraphs[-1]
docPr = p._p.xpath('.//wp:docPr')[0]
docPr.set('descr','Scatter plot comparing correctness percentage and mean response time across six RAG configurations.')
p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2)
doc.add_paragraph('Figure 1. Correctness–latency tradeoff across six RAG configurations.',style='Caption')
heading('4. Challenges and Responses',1)
add_p('An early evaluation run returned HTTP 404 errors for every generation request while embedding and retrieval continued. The artifacts therefore showed near-zero latency and zero quality—numbers that could have been mistaken for poor model performance. The evaluator was changed to exit unsuccessfully when model requests fail and to preserve failed-run artifacts for diagnosis without treating them as quality evidence.')
add_p('A second challenge was that more context was not consistently better. Increasing top-k expanded the prompt and slowed generation without improving retrieval hits. Conversely, top-k one preserved strong answer quality while reducing latency and retrieved context. The response was to evaluate configurations with a combined quality–speed score and retain the full metric set so the selection remained auditable.')

# PAGE 6 — discussion
page_break(); heading('5. Discussion',1)
add_p('The technical evidence supports retrieval-augmented generation for source-specific educational assistance. Grounding substantially improved correctness and boundary handling, and the selected configuration showed that a small local model can respond in roughly one second on the evaluated hardware. These findings strengthen the case for a controlled assistant that prioritizes instructor-approved material over unconstrained general knowledge.')
add_p('The results also show why system quality cannot be represented by a single accuracy number. A deployment can fail because the model service is unavailable, retrieval misses the necessary evidence, generation omits part of the evidence, or an answer adds unsupported information. Archivist’s evaluation therefore separates operational availability, retrieval, correctness, groundedness, hallucination handling, and latency. This decomposition makes debugging more useful and reduces the risk of interpreting infrastructure errors as model behavior.')
add_p('However, the project has not yet demonstrated that guided workflows significantly reduce setup difficulty for real administrators. The MVP makes that study possible, but the current experiments measure answer behavior rather than human task performance. The 15-question dataset is too small for broad generalization, the questions were researcher-written, the source covers one short section, and one manual reviewer scored the final baseline. The optimization quality scores are an AI-assisted draft pending final human verification. No statistical significance is claimed.')
add_p('The application also retains engineering limitations. The current extraction path supports text-based PDFs but not scanned or image-only documents. Embeddings are stored as JSON in SQLite and scanned in process, favoring inspectability over large-scale retrieval. Reindexing has route and interface foundations but not a fully observable background-job worker. Chat history is grouped by workspace rather than named conversations. Production hardening also requires CSRF protection, rate limiting, production TLS, richer account administration, and broader operational monitoring.')
add_p('Finally, local deployment reduces dependence on hosted APIs but does not eliminate governance responsibilities. Schools still need clear policies for approved sources, account access, retention, model updates, and review of high-impact answers. Citations help students and instructors verify responses, but they should support—not replace—source reading and instructional judgment. The system should communicate uncertainty and refuse unsupported questions rather than presenting fluent speculation as course knowledge.')

# PAGE 7 — future, conclusion, references
page_break(); heading('6. Future Work',1)
add_p('The immediate next phase is a stakeholder usability study. Participants should complete first-run setup, create a workspace, upload a source, create and assign a student, and obtain a cited answer. Measures should include task completion rate, time on task, setup errors, recovery behavior, requests for assistance, and perceived confidence before and after the workflow. A comparison against an unguided setup path would directly test the primary research question.')
add_p('The evidence base should also expand to include a larger independently authored and held-out question set, multiple source types and academic domains, repeated runs across seeds, multiple human reviewers, and inter-rater agreement. Targeted experiments should evaluate token-aware chunking, reranking, metadata filters, alternative embeddings, larger local chat models, and answer-level citation precision. Regression cases should preserve every observed retrieval, generation, and availability failure.')
add_p('In parallel, product work should add OCR and layout-aware extraction, observable background ingestion and reindexing jobs, named conversation sessions, richer source previews, model-health guidance, secure password administration, CSRF protection, rate limiting, TLS, and scalable vector storage. Offline Linux deployment should continue to verify network isolation and provide reversible, plain-language installation guidance.')
heading('7. Conclusion',1)
add_p('Archivist demonstrates a credible technical path for private, source-grounded educational AI. Retrieval raised correctness by 30 percentage points in the final baseline, while configuration tuning produced a faster selected system with 96.7% correctness and strong groundedness. The project’s most important contribution is the integration of local inference, trusted-source retrieval, guided administration, reproducible evaluation, and explicit failure analysis into one deployable MVP. The next milestone is not another isolated model score; it is evidence that real administrators can successfully complete the workflow with fewer errors and greater confidence.')
heading('References and Project Artifacts',2)
refs=[
    '[1] Archivist repository, README.md and docs/capstone-notes.md, project scope, architecture, limitations, and roadmap.',
    '[2] results/baseline_results.csv and results/report_summary.txt, final 15-question no-RAG versus RAG manual evaluation.',
    '[3] results/optimization_summary.csv, results/optimization_metadata.json, and results/optimization_report_summary.txt, six-configuration RAG study.',
    '[4] evaluation/README.md and scripts/evaluate.py, versioned evaluation protocol and reproducible runner.',
    '[5] docs/architecture.md, docs/setup.md, and docs/linux-deployment.md, implementation and deployment design.'
]
for ref in refs:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2); p.paragraph_format.first_line_indent=Inches(-0.2); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(ref); set_font(r,9,False,False,MUTED)

# Core metadata and accessibility
doc.core_properties.title='Archivist: Final Technical Report'
doc.core_properties.subject='MSAI 699 Capstone Project'
doc.core_properties.author='Nick Moore'
doc.core_properties.keywords='local AI, education, retrieval-augmented generation, RAG, guided workflows'
doc.save(OUT)
print(OUT)
