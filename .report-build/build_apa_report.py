from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path('/Users/nickmoore/Desktop/school/summer_2026/msai_699/archivist')
OUT = ROOT / 'Archivist_Final_Technical_Report.docx'
FONT = 'Times New Roman'

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = Inches(0.5)
sec.footer_distance = Inches(0.5)

def font(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn('w:ascii'), FONT)
    rpr.rFonts.set(qn('w:hAnsi'), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '
    separate = OxmlElement('w:fldChar'); separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t'); text.text = '1'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    for el in (begin, instr, separate, text, end): run._r.append(el)
    font(run)

page_number(sec.header.paragraphs[0])

# APA 7 student-paper style sheet.
normal = doc.styles['Normal']
normal.font.name = FONT; normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 2
normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(0)
for name in ('Heading 1', 'Heading 2', 'Heading 3'):
    st = doc.styles[name]
    st.font.name = FONT; st.font.size = Pt(12); st.font.bold = True
    st.font.color.rgb = RGBColor(0, 0, 0)
    st.paragraph_format.line_spacing = 2
    st.paragraph_format.space_before = Pt(0); st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.keep_with_next = True
doc.styles['Heading 1'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.styles['Heading 2'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.styles['Heading 3'].font.italic = True

def para(text='', indent=True, align=None, bold=False, italic=False, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    if indent: p.paragraph_format.first_line_indent = Inches(0.5)
    if align is not None: p.alignment = align
    p.paragraph_format.keep_with_next = keep
    r = p.add_run(text); font(r, bold=bold, italic=italic)
    return p

def h1(text):
    p = doc.add_paragraph(style='Heading 1'); p.paragraph_format.line_spacing = 2
    r = p.add_run(text); font(r, bold=True)
    return p

def title_line(text, bold=False):
    return para(text, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, bold=bold)

def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tcPr = cell._tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for side, val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        e = tcMar.find(qn(f'w:{side}'))
        if e is None: e = OxmlElement(f'w:{side}'); tcMar.append(e)
        e.set(qn('w:w'), str(val)); e.set(qn('w:type'), 'dxa')

def table_borders(table):
    tblPr = table._tbl.tblPr
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None: tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','bottom','insideH'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'8'); e.set(qn('w:color'),'000000'); borders.append(e)
    for edge in ('left','right','insideV'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'),'nil'); borders.append(e)
    tblPr.append(borders)

def results_table():
    para('Table 1', indent=False, bold=True, keep=True)
    para('Baseline No-RAG and RAG Results', indent=False, italic=True, keep=True)
    headers = ['Metric','No RAG','RAG','Difference']
    rows = [
        ('Correctness','63.3%','93.3%','+30.0 points'),
        ('Mean response time','1.11 s','1.25 s','+0.14 s'),
        ('Retrieval Hit Rate@3','—','100%','—'),
        ('Average groundedness','—','0.93','—'),
        ('Hallucination handling','—','100%','—'),
    ]
    t = doc.add_table(rows=1, cols=4); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    trPr = t.rows[0]._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'), 'true'); trPr.append(tblHeader)
    widths=[Inches(2.5),Inches(1.2),Inches(1.2),Inches(1.6)]
    for i,h in enumerate(headers):
        cell=t.cell(0,i); cell.width=widths[i]; set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing=1; p.paragraph_format.space_after=Pt(0)
        r=p.add_run(h); font(r,10,bold=True)
    for row in rows:
        cells=t.add_row().cells
        for i,value in enumerate(row):
            cells[i].width=widths[i]; set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT if i==0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing=1; p.paragraph_format.space_after=Pt(0)
            r=p.add_run(value); font(r,10)
    table_borders(t)
    para('Note. Results are from the 15-question project evaluation (Moore, 2026).', indent=False, italic=True)

# Title page: APA student paper, using only known metadata.
for _ in range(4): para('', indent=False)
title_line('Archivist: Investigating Guided Workflows for Reducing Technical Barriers to Local Artificial Intelligence Adoption', bold=True)
title_line('Nick Moore')
title_line('MSAI 699: Capstone Project')
title_line('August 12, 2026')
doc.add_page_break()

# Body
title_line('Archivist: Investigating Guided Workflows for Reducing Technical Barriers to Local Artificial Intelligence Adoption', bold=True)

para('Educational institutions want the accessibility of conversational artificial intelligence (AI) without surrendering control over course content, student interactions, and institutional data. Hosted assistants lower the barrier to experimentation, but they can introduce concerns involving privacy, source authority, recurring cost, network dependence, and unsupported answers. Local models keep inference and storage under institutional control, but setup and maintenance commonly require familiarity with containers, model servers, embeddings, retrieval configuration, and networking.')
para('This project investigates whether guided workflows can reduce the technical barriers associated with configuring and deploying local AI assistants for educational use. Archivist makes the question concrete through an end-to-end minimum viable product (MVP). An administrator can create a workspace, upload approved sources, add and assign students, and provide a grounded question-answering experience. The work had three objectives: build a complete local application, test whether retrieval improves source-specific answers, and tune retrieval for a practical balance between quality and response time.')
para('Archivist uses a Go HTTP server, server-rendered HTML, SQLite storage, and Ollama for local chat and embedding inference (Moore, 2026). Uploaded files are extracted, divided into overlapping chunks, embedded, and stored with pipeline-version metadata. For each student question, the retrieval service embeds the query, ranks stored chunks using cosine similarity, supplies selected evidence to a constrained prompt, and returns an answer with source references. Docker Compose packages the application and model service into a repeatable local deployment.')

h1('Methodology')
para('Development followed an iterative workflow: define the stakeholder path, implement the smallest complete vertical slice, test it locally, document failure modes, and convert those failures into regression checks. The resulting path begins with first-run administration and continues through workspace creation, source ingestion, student assignment, retrieval, and cited response generation. The web and terminal interfaces share application services so business rules are not duplicated.')
para('The baseline experiment used 15 versioned questions about an instructor-prepared paraphrase of an OpenStax data-science section (Ault et al., 2025). Questions covered definitions, comparisons, recall, reasoning, synthesis, and two deliberately unanswerable items. Both conditions used the same local chat model and deterministic generation settings. The no-RAG condition received only the question. The retrieval-augmented generation (RAG) condition also received the three most similar source chunks, using 1,000-character chunks, 200-character overlap, nomic-embed-text embeddings, and cosine similarity.')
para('Each run captured responses, retrieved chunk identifiers, similarity scores, retrieval time, generation time, and total response time. Manual review scored correctness on a 0–2 rubric: 2 for correct and complete, 1 for partially correct or incomplete, and 0 for incorrect or missing. RAG answers were also scored for groundedness against the retrieved text. Retrieval hit rate measured whether the selected context contained the evidence required by the reference answer. Hallucination handling measured whether the model avoided unsupported content when the source could not answer the question.')
para('A second experiment evaluated six RAG configurations with the same 15 questions. It varied chunk size and overlap—600/100, 800/150, 1,000/200, and 1,400/250 characters—and varied top-k at 1, 3, and 5 for the 1,000/200 setting. A project-specific selection score combined correctness, groundedness, retrieval success, and response time. The dataset, source, settings, run metadata, raw answers, review files, and summaries were retained for auditability (Moore, 2026).')

h1('Results')
para('Direct generation achieved 63.3% of the maximum correctness score. Adding retrieved course context increased correctness to 93.3%, a gain of 30.0 percentage points. Retrieval Hit Rate@3 was 100%, average groundedness was 0.93 on a 0–1 scale, and hallucination handling was 100%. Mean response time increased modestly from 1.11 seconds without RAG to 1.25 seconds with RAG (Moore, 2026).')
results_table()
para('The largest difference appeared when questions required source-specific wording or boundaries. Direct generation could produce plausible general knowledge while missing the assigned material’s intended explanation. On an out-of-scope question about transformer embeddings, the no-RAG answer invented a plausible explanation even though the section did not contain one; RAG correctly declined to answer from the available material. Retrieval success did not guarantee complete generation, however. Some answers used the correct evidence but omitted a qualification or overstated the source, which justified measuring retrieval, groundedness, and correctness separately.')
para('Across the optimization study, correctness ranged from 93.3% to 100%, while mean response time ranged from 1.07 to 1.88 seconds. The selected top-k-one configuration used 1,000-character chunks with 200-character overlap and retrieved one chunk. It achieved 96.7% correctness, 0.97 groundedness, and a 1.07-second mean response time. Relative to the top-k-three baseline, correctness increased 3.3 points and mean latency fell 0.81 seconds, or 42.9% (Moore, 2026).')
para('Figure 1', indent=False, bold=True, keep=True)
para('Correctness–Latency Tradeoff Across Six RAG Configurations', indent=False, italic=True, keep=True)
pic = doc.add_picture(str(ROOT/'results/charts/correctness_latency_tradeoff.png'), width=Inches(4.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
docPr = doc.paragraphs[-1]._p.xpath('.//wp:docPr')[0]
docPr.set('descr','Scatter plot comparing correctness percentage and mean response time for six retrieval configurations.')
para('Note. The top-k-one configuration was selected for its combined quality and latency performance (Moore, 2026).', indent=False, italic=True)

h1('Discussion')
para('The technical evidence supports RAG for source-specific educational assistance. Grounding substantially improved correctness and boundary handling, and the selected configuration showed that a small local model can respond in approximately one second on the evaluated hardware. The findings strengthen the case for an assistant that prioritizes instructor-approved material over unconstrained general knowledge.')
para('The project also encountered two consequential challenges. An early evaluation run returned HTTP 404 errors for every generation request while embedding and retrieval continued. Its artifacts showed near-zero latency and zero quality, which could have been mistaken for model performance. The evaluator was changed to fail when model requests fail while preserving artifacts for diagnosis. In addition, more context was not consistently better. Increasing top-k slowed generation without improving retrieval hits; top-k one preserved strong quality while reducing latency and prompt size.')
para('System quality therefore cannot be represented by a single accuracy number. A deployment can fail because the model service is unavailable, retrieval misses evidence, generation omits part of the evidence, or an answer adds unsupported information. Archivist separates operational availability, retrieval, correctness, groundedness, hallucination handling, and latency so failures can be interpreted and corrected appropriately.')
para('The study does not yet demonstrate that guided workflows reduce setup difficulty for real administrators. The current experiments measure answer behavior rather than human task performance. The question set is small and researcher-written, the source covers one short section, and one manual reviewer scored the final baseline. The optimization scores are an AI-assisted draft pending final human verification; no statistical significance is claimed.')
para('The application also retains engineering and governance limitations. It supports text-based PDFs but not scanned documents. Embeddings are stored as JSON in SQLite and scanned in process, favoring inspectability over scale. Reindexing lacks a fully observable background-job worker, and production hardening requires CSRF protection, rate limiting, TLS, richer account administration, and monitoring. Local deployment reduces dependence on hosted APIs but does not remove the need for policies governing approved sources, access, retention, model updates, and review of high-impact answers.')

h1('Future Work')
para('The immediate next phase should evaluate the guided workflow with administrators. Participants should complete first-run setup, create a workspace, upload a source, create and assign a student, and obtain a cited answer. Measures should include completion rate, time on task, setup errors, recovery behavior, assistance requests, and perceived confidence. Comparing this experience with an unguided setup path would directly test the project’s primary research question.')
para('The evidence base should also expand to an independently authored, held-out question set; multiple sources and academic domains; repeated runs across seeds; multiple reviewers; and interrater agreement. Targeted experiments should examine token-aware chunking, reranking, metadata filters, alternative embeddings, larger local models, and citation precision. Product work should add optical character recognition, observable background ingestion, named conversations, richer source previews, model-health guidance, stronger security controls, and scalable vector storage.')

h1('Conclusion')
para('Archivist demonstrates a credible technical path for private, source-grounded educational AI. Retrieval raised correctness by 30 percentage points in the final baseline, while configuration tuning produced a faster selected system with 96.7% correctness and strong groundedness. The project integrates local inference, trusted-source retrieval, guided administration, reproducible evaluation, and explicit failure analysis into one deployable MVP. The next milestone is evidence that real administrators can complete the workflow with fewer errors and greater confidence.')

doc.add_page_break()
h1('References')
refs = [
    ('Ault, S. V., Liao, S. N., & Musolino, L. (2025). ', 'Principles of data science', '. OpenStax. https://openstax.org/details/books/principles-data-science'),
    ('Moore, N. (2026). ', 'Archivist', ' [Computer software and unpublished project data].'),
]
for lead, title, tail in refs:
    p=doc.add_paragraph(); p.paragraph_format.line_spacing=2; p.paragraph_format.space_after=Pt(0)
    p.paragraph_format.left_indent=Inches(0.5); p.paragraph_format.first_line_indent=Inches(-0.5)
    r=p.add_run(lead); font(r)
    r=p.add_run(title); font(r, italic=True)
    r=p.add_run(tail); font(r)

doc.core_properties.title='Archivist: Investigating Guided Workflows for Reducing Technical Barriers to Local Artificial Intelligence Adoption'
doc.core_properties.author='Nick Moore'
doc.core_properties.subject='MSAI 699 Capstone Project'
doc.save(OUT)
print(OUT)
